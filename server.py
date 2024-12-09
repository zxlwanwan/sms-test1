from docx import Document
from docx.shared import Pt
import docx

# 创建文档
doc = Document()

# 添加标题
doc.add_heading('智能电动床功能说明书', 0)

# 添加产品信息段落
para = doc.add_paragraph()
para.add_run('产品名称: ').bold = True
para.add_run('智能电动床')

# 添加功能表格
doc.add_heading('功能模块', level=1)
table = doc.add_table(rows=3, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '功能'
hdr_cells[1].text = '参数'
hdr_cells[2].text = '描述'

row_cells = table.rows[1].cells
row_cells[0].text = '智能调节'
row_cells[1].text = '多角度调节'
row_cells[2].text = '根据睡姿提供不同角度的自动调节'

row_cells = table.rows[2].cells
row_cells[0].text = '睡眠监测'
row_cells[1].text = '健康监测'
row_cells[2].text = '记录并分析睡眠质量'

# 插入图片
doc.add_paragraph('产品图片:')
doc.add_picture('bed_image.jpg', width=docx.shared.Inches(3))

# 添加页脚
section = doc.sections[0]
footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = "智能电动床说明书"

# 保存文件
doc.save('智能电动床功能说明书.docx')
